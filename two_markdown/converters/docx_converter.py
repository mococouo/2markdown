from __future__ import annotations

import mimetypes
import zipfile
from pathlib import Path

from two_markdown.context import ConversionContext
from two_markdown.errors import MissingDependencyError
from two_markdown.models import ConvertedDocument
from two_markdown.utils import markdown_table


class DocxConverter:
    name = "docx"
    extensions = {".docx"}

    def convert(self, path: Path, context: ConversionContext) -> ConvertedDocument:
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except Exception as exc:
            raise MissingDependencyError("python-docx", "DOCX conversion") from exc

        document = Document(path)
        parts = [f"# {path.stem}"]
        warnings: list[str] = []
        used_media: set[str] = set()

        for paragraph in document.paragraphs:
            chunks = _paragraph_to_markdown(paragraph, context, qn, used_media, warnings)
            text = "".join(chunks).strip()
            if not text:
                continue
            style_name = (paragraph.style.name if paragraph.style else "").lower()
            if style_name.startswith("heading"):
                level = _heading_level(style_name)
                parts.append(f"{'#' * level} {text}")
            elif "list bullet" in style_name:
                parts.append(f"- {text}")
            elif "list number" in style_name:
                parts.append(f"1. {text}")
            else:
                parts.append(text)

        for table in document.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                parts.append(markdown_table(rows))

        extra_media = _extract_docx_media(path)
        appended = 0
        for filename, data in extra_media:
            if filename in used_media:
                continue
            appended += 1
            link = context.add_attachment(filename, data, alt_text=Path(filename).stem)
            parts.append(link)
        if appended:
            warnings.append(f"{appended} image(s) were not placed inline and appended at the end.")

        metadata = _core_properties(document) if context.options.include_metadata else {}

        return ConvertedDocument(
            title=metadata.get("title") or path.stem,
            markdown="\n\n".join(parts),
            converter=self.name,
            attachments=context.attachments,
            warnings=warnings,
            metadata=metadata,
        )


def _paragraph_to_markdown(
    paragraph: object,
    context: ConversionContext,
    qn: object,
    used_media: set[str],
    warnings: list[str],
) -> list[str]:
    out: list[str] = []
    try:
        element = paragraph._p  # type: ignore[attr-defined]
        part = paragraph.part  # type: ignore[attr-defined]
    except Exception:
        out.append(paragraph.text)  # type: ignore[attr-defined]
        return out

    for child in element.iterchildren():
        tag = child.tag
        if tag == qn("w:r"):
            out.append(_run_to_markdown(child, qn, part, context, used_media, warnings))
        elif tag == qn("w:hyperlink"):
            out.append(_hyperlink_to_markdown(child, qn, part, context))
    return out


def _run_to_markdown(
    run_element: object,
    qn: object,
    part: object,
    context: ConversionContext,
    used_media: set[str],
    warnings: list[str],
) -> str:
    image_token = _run_image_token(run_element, qn, part, context, used_media, warnings)
    if image_token:
        return image_token

    bold = _element_bool(run_element, qn("w:b"))
    italic = _element_bool(run_element, qn("w:i"))
    underline = _element_bool(run_element, qn("w:u"))
    text = _run_text(run_element, qn)
    if not text:
        return ""
    wrapped = text
    if bold:
        wrapped = f"**{wrapped}**"
    if italic:
        wrapped = f"*{wrapped}*"
    if underline:
        wrapped = f"<u>{wrapped}</u>"
    return wrapped


def _run_text(run_element: object, qn: object) -> str:
    pieces: list[str] = []
    for node in run_element.iter():  # type: ignore[attr-defined]
        if node.tag == qn("w:t"):
            pieces.append(node.text or "")
        elif node.tag == qn("w:tab"):
            pieces.append("\t")
        elif node.tag == qn("w:br") or node.tag == qn("w:cr"):
            pieces.append("\n")
    return "".join(pieces)


def _element_bool(element: object, tag: object) -> bool:
    found = list(element.iter(tag))  # type: ignore[attr-defined]
    if not found:
        return False
    value = found[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
    return value not in {"0", "false"}


def _run_image_token(
    run_element: object,
    qn: object,
    part: object,
    context: ConversionContext,
    used_media: set[str],
    warnings: list[str],
) -> str:
    embed_attr = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    for blip in run_element.iter("{http://schemas.openxmlformats.org/drawingml/2006/main}blip"):  # type: ignore[attr-defined]
        rid = blip.get(embed_attr)
        if not rid:
            continue
        try:
            rel = part.rels[rid]  # type: ignore[attr-defined]
            target = rel.target_part
            blob = target.blob
            partname = Path(str(target.partname)).name
        except Exception as exc:
            warnings.append(f"Could not extract inline image: {exc}")
            continue
        used_media.add(partname)
        extension = Path(partname).suffix or mimetypes.guess_extension(getattr(target, "content_type", "")) or ".png"
        if not Path(partname).suffix:
            partname = f"{partname}{extension}"
        return context.add_attachment(partname, blob, alt_text=Path(partname).stem)
    return ""


def _hyperlink_to_markdown(hyperlink_element: object, qn: object, part: object) -> str:
    rid = hyperlink_element.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")  # type: ignore[attr-defined]
    text_parts: list[str] = []
    for run_element in hyperlink_element.iterchildren(qn("w:r")):  # type: ignore[attr-defined]
        text_parts.append(_run_text(run_element, qn))
    text = "".join(text_parts).strip()
    url = ""
    if rid:
        try:
            url = part.rels[rid].target_ref  # type: ignore[attr-defined]
        except Exception:
            url = ""
    if not text:
        return ""
    if url:
        return f"[{text}]({url})"
    return text


def _heading_level(style_name: str) -> int:
    for token in reversed(style_name.split()):
        if token.isdigit():
            return max(1, min(int(token), 6))
    return 2


def _core_properties(document: object) -> dict[str, str]:
    metadata: dict[str, str] = {}
    try:
        props = document.core_properties  # type: ignore[attr-defined]
    except Exception:
        return metadata
    pairs = {
        "author": getattr(props, "author", None),
        "subject": getattr(props, "subject", None),
        "keywords": getattr(props, "keywords", None),
        "category": getattr(props, "category", None),
        "comments": getattr(props, "comments", None),
    }
    for key, value in pairs.items():
        if value and str(value).strip():
            metadata[key] = str(value).strip()
    title = getattr(props, "title", None)
    if title and str(title).strip():
        metadata["title"] = str(title).strip()
    created = getattr(props, "created", None)
    if created:
        metadata["created"] = _iso_dt(created)
    modified = getattr(props, "modified", None)
    if modified:
        metadata["modified"] = _iso_dt(modified)
    return metadata


def _iso_dt(value: object) -> str:
    iso = getattr(value, "isoformat", None)
    if callable(iso):
        try:
            return iso()
        except Exception:
            pass
    return str(value)


def _extract_docx_media(path: Path) -> list[tuple[str, bytes]]:
    media: list[tuple[str, bytes]] = []
    with zipfile.ZipFile(path) as archive:
        for info in archive.infolist():
            if not info.filename.startswith("word/media/"):
                continue
            filename = Path(info.filename).name
            if not Path(filename).suffix:
                guessed = mimetypes.guess_extension(info.filename)
                filename = f"{filename}{guessed or '.bin'}"
            media.append((filename, archive.read(info.filename)))
    return media
